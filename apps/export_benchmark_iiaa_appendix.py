from __future__ import annotations

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = Path("debug_dump/final_benchmark_iiaa_results.jsonl")
OUT = Path("debug_dump/Appendice_Benchmark_IIAA.docx")


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size in [("Title", 16), ("Heading 1", 14), ("Heading 2", 12)]:
        st = styles[style_name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string("1F4E79")


def add_paragraph(doc: Document, text: str, style: str | None = None, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic
    return p


def add_case_summary_table(doc: Document, row: dict) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valore"
    for c in hdr:
        set_cell_shading(c, "D9EAF7")
        set_cell_margins(c, 100, 100, 100, 100)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
                r.bold = True

    qp = row.get("query_plan") or {}
    fields = [
        ("Case ID", row.get("case_id")),
        ("Label", row.get("label")),
        ("Status", row.get("status")),
        ("Question type", qp.get("question_type")),
        ("Source preference", qp.get("source_preference")),
        ("Target standards", ", ".join(qp.get("target_standards") or [])),
        ("Query length original", row.get("query_len_original")),
        ("Query length embedded", row.get("query_len_embedded")),
        ("Query truncated", row.get("query_was_truncated")),
        ("Retrieval strategy", row.get("retrieval_query_strategy")),
        ("Used / candidate citations", f"{len(row.get('used_citations') or [])}/{len(row.get('citations') or [])}"),
        ("Error", row.get("error") or ""),
    ]

    for k, v in fields:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = "" if v is None else str(v)
        for c in cells:
            set_cell_margins(c, 90, 90, 90, 90)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)

    doc.add_paragraph("")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(f"Missing source file: {SRC}")

    rows = [json.loads(line) for line in SRC.read_text(encoding="utf-8").splitlines() if line.strip()]

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Appendice tecnica\nRisultati del benchmark finale IIAA")
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Project work – benchmark locale grounded IAS/IFRS")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = True

    doc.add_paragraph("")

    add_paragraph(
        doc,
        "La presente appendice riporta gli esiti del benchmark finale del prototipo IIAA sui prompt canonici utilizzati per il confronto con i chatbot generalisti. Per ciascun caso sono riportati identificativo, classificazione della domanda, strategia di retrieval embedding-based, stato di esecuzione, rapporto tra citazioni usate e citazioni candidate e risposta prodotta dal sistema.",
    )

    add_paragraph(
        doc,
        "Il benchmark finale è stato eseguito in ambiente locale, con corpus controllato e pipeline grounded. I risultati qui riportati hanno quindi funzione di documentazione tecnica, supporto valutativo e base per la discussione metodologica del Capitolo 4.",
    )

    ok_count = sum(1 for r in rows if r.get("status") == "ok")
    err_count = sum(1 for r in rows if r.get("status") == "error")

    add_paragraph(doc, f"Casi totali: {len(rows)}. Casi completati con successo: {ok_count}. Casi in errore: {err_count}.", bold=True)
    doc.add_paragraph("")

    for idx, row in enumerate(rows, start=1):
        add_paragraph(doc, f"Caso {idx} – {row.get('case_id','')}", style="Heading 1")
        add_case_summary_table(doc, row)

        preview = (row.get("embedding_query_preview") or row.get("retrieval_query_preview") or "").strip()
        if preview:
            add_paragraph(doc, "Query embedding-based utilizzata", style="Heading 2")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(preview)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10.5)
            r.italic = True

        answer = (row.get("answer") or "").strip()
        add_paragraph(doc, "Risposta del sistema", style="Heading 2")
        if answer:
            for chunk in answer.split("\n"):
                chunk = chunk.strip()
                if chunk:
                    add_paragraph(doc, chunk)
        else:
            add_paragraph(doc, "(nessuna risposta disponibile)")

        used = [c.get("cite_key") for c in (row.get("used_citations") or []) if c.get("cite_key")]
        cand = [c.get("cite_key") for c in (row.get("citations") or []) if c.get("cite_key")]

        add_paragraph(doc, "Citazioni usate", style="Heading 2")
        add_paragraph(doc, ", ".join(used) if used else "(nessuna)")

        add_paragraph(doc, "Citazioni candidate", style="Heading 2")
        add_paragraph(doc, ", ".join(cand) if cand else "(nessuna)")

        if idx < len(rows):
            doc.add_page_break()

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved {OUT}")


if __name__ == "__main__":
    main()
